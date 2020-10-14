import sys
import os
import re
import numpy as np
import cv2
import docx

from docx import shared,Document
import time
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None
def cv_imread(file_path):
    cv_img = cv2.imdecode(np.fromfile(file_path,dtype=np.uint8),-1)
    return cv_img
if __name__ == '__main__':
    c_lsit = os.listdir()
    c_lsit.remove(".idea")
    #获取学号
    # d_list = [f[:10] for f in c_lsit ]
    # print(d_list[6])
    # d ="2016210035陈瑾1"

    mo = u"[\u4e00-\u9fa5]+"
    e_list = []
    f_list = []
    for i in c_lsit:
        idc = i[:10]
        # print(idc)#str
        # print(type(idc))
        r = re.findall(mo, i)

        if len(r)!=0:
            print(idc+r[0])
            # print(r[0])
            e_list.append(r[0])
    # print(e_list[10])

    for x in c_lsit:
        #处理图片
        if "jpg" in x:


            ##在图片里面匹配学号+名字
            idc = x[:10]#获取学号
            # print(idc)#str
            # print(type(idc))
            r = re.findall(mo, x)

            if len(r) != 0:
                print(idc + r[0])



                # print(r[0])
                e_list.append(r[0])
            ##
            image = cv_imread(x)
            pic = cv2.resize(image, (800, 750), interpolation=cv2.INTER_CUBIC)

            cv2.imshow(x,pic)
            cv2.moveWindow(x, 250, 50)

            # print(pic.shape)
            # cv2.resizeWindow(x,400,400)

            recode = cv2.waitKey(0)
            if recode==32:
                cv2.imencode('.jpg', pic)[1].tofile("operated\\"+x)
                if (os.path.exists(idc + r[0] +"end" ".docx")):
                    doc = Document(idc + r[0] +"end" ".docx")
                    # doc.save(idc + r[0] + ".docx")

                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] +"end" ".docx")
                else:
                    doc = docx.Document("demo.docx")
                    delete_paragraph(doc.paragraphs[0])
                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save("operated\\"+idc + r[0] + "end" ".docx")
            elif recode == 97:
                img = np.rot90(pic, 1)
                cv2.imencode('.jpg', pic)[1].tofile("operated\\"+x)
                #如果存在
                if (os.path.exists(idc + r[0] +"end" ".docx")):
                    doc = Document(idc + r[0] +"end" ".docx")
                    # doc.save(idc + r[0] + ".docx")

                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] +"end" ".docx")
                else:
                    doc = docx.Document("demo.docx")
                    delete_paragraph(doc.paragraphs[0])
                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] + "end" ".docx")

            elif recode == 100:
                img = np.rot90(pic,-1)
                cv2.imencode('.jpg', pic)[1].tofile("operated\\"+x)
                if (os.path.exists(idc + r[0] + "end" ".docx")):
                    doc = Document(idc + r[0] + "end" ".docx")
                    # doc.save(idc + r[0] + ".docx")

                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] + "end" ".docx")
                else:
                    doc = docx.Document("demo.docx")
                    delete_paragraph(doc.paragraphs[0])
                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] + "end" ".docx")
            elif recode == 119:
                img = np.rot90(pic,-2)
                cv2.imencode('.jpg', pic)[1].tofile("operated\\"+x)
                if (os.path.exists(idc + r[0] + "end" ".docx")):
                    doc = Document(idc + r[0] + "end" ".docx")
                    # doc.save(idc + r[0] + ".docx")

                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] + "end" ".docx")
                else:
                    doc = docx.Document("demo.docx")
                    delete_paragraph(doc.paragraphs[0])
                    doc.add_picture(x, width=shared.Inches(7))
                    doc.save(idc + r[0] + "end" ".docx")
            else:
                print("出现错误请停止处理！！！")


            print(recode)
            cv2.destroyAllWindows()
            print("处理成功\n")

            ###


