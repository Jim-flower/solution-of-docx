import sys
import os
import re
import numpy as np
import cv2
import docx

from docx import shared,Document
import time
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None
def cv_imread(file_path):
    cv_img = cv2.imdecode(np.fromfile(file_path,dtype=np.uint8),-1)
    return cv_img
def get_pictures(word_path, result_path):
    """
    提取word文档内的图片
    :param word_path: word文件
    :param result_path: 结果目录
    :return:
    """
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            if not os.path.exists(result_path):
                os.makedirs(result_path)
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            # print(os.sep)
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = f'{new_name}_{img_name}'
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)
if __name__ == '__main__':
    c_lsit = os.listdir()
    # print(c_lsit)
    c_lsit.remove(".idea")
    #获取学号
    # d_list = [f[:10] for f in c_lsit ]
    # print(d_list[6])
    # d ="2016210035陈瑾1"

    mo = u"[\u4e00-\u9fa5]+"
    e_list = []
    f_list = []

    for x in c_lsit:
        #处理图片
        if "png" in x:

            print("进来")
            ##在图片里面匹配学号+名字
            idc = x[:10]#获取学号
            # print(idc)#str
            # print(type(idc))
            r = re.findall(mo, x)

            if len(r) != 0:
                print(idc + r[0])



                # print(r[0])
                e_list.append(r[0])
            ##
            image = cv_imread(x)
            pic = cv2.resize(image, (800, 750), interpolation=cv2.INTER_CUBIC)

            cv2.imshow(x,pic)
            cv2.moveWindow(x, 250, 50)

            # print(pic.shape)
            # cv2.resizeWindow(x,400,400)

            recode = cv2.waitKey(0)
            cv2.destroyAllWindows()
            print(r)
            if len(r)!=0:
                path = "operated\\" + idc + r[0] + "end" ".docx"  # 学号+名字+“end"
                path1 = "operated\\" + x
                if recode == 32:
                    # 不做任何改变 保存在op文件夹
                    cv2.imencode('.png', pic)[1].tofile("operated\\" + x)
                elif recode == 97:
                    img = np.rot90(pic, 1)
                    cv2.imencode('.png', img)[1].tofile("operated\\" + x)
                elif recode == 100:

                    img = np.rot90(pic, -1)
                    cv2.imencode('.png', img)[1].tofile("operated\\" + x)
                elif recode == 119:
                    img = np.rot90(pic, -2)
                    cv2.imencode('.png', img)[1].tofile("operated\\" + x)

                os.remove(x)
                if (os.path.exists(path)):
                    doc = Document(path)  # 打开文件
                    doc.add_picture(path1, width=shared.Inches(7))  # 插入图片
                    doc.save(path)
                else:
                    doc = Document("demo.docx")
                    delete_paragraph(doc.paragraphs[0])
                    doc.add_picture(path1, width=shared.Inches(7))
                    doc.save(path)

                os.remove(path1)





