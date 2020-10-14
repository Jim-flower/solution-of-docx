import docx

from docx import shared
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

if __name__ == '__main__':
    doc = docx.Document("demo.docx")
   #先删除第一行
    delete_paragraph(doc.paragraphs[0])
    for i in doc.paragraphs:
        print(i)
    # doc.add_heading("")
    # doc.add_paragraph('python-docx 基础讲解（二）')
    #插入图片
    doc.add_picture('2017210220 高盛(1).jpg', width=shared.Inches(7))
    doc.add_picture('2017210220 高盛(1).jpg', width=shared.Inches(7))
    doc.add_picture('2017210220 高盛(1).jpg', width=shared.Inches(7))

    doc.save("demo1.docx")

